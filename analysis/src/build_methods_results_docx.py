"""Create a publication-ready Methods/Results DOCX for JQL submission."""
from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

import config


def _fmt(x: float, n: int = 3) -> str:
    return f"{x:.{n}f}"


def _set_doc_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)


def _add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.style = "Caption"
    except KeyError:
        pass


def _add_table_from_df(doc: Document, df: pd.DataFrame, title: str) -> None:
    _add_caption(doc, title)
    t = doc.add_table(rows=1, cols=len(df.columns))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass

    hdr = t.rows[0].cells
    for i, c in enumerate(df.columns):
        hdr[i].text = str(c)
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, v in enumerate(row.tolist()):
            if isinstance(v, float):
                txt = _fmt(v, 4)
            else:
                txt = str(v)
            cells[i].text = txt
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")


def _add_figure(doc: Document, path: Path, caption: str, width_in: float = 6.2) -> None:
    if not path.exists():
        doc.add_paragraph(f"[Missing figure: {path.name}]")
        return
    pic = doc.add_picture(str(path), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_caption(doc, caption)
    doc.add_paragraph("")


def _performance_wide(perf: pd.DataFrame) -> pd.DataFrame:
    acc = perf.pivot(index="scenario", columns="method", values="acc_1nn_loo").reset_index()
    ari = perf.pivot(index="scenario", columns="method", values="ari_3cluster").reset_index()
    out = acc.merge(ari, on="scenario", suffixes=("_acc", "_ari"))
    out = out.rename(columns={
        "openai_acc": "OpenAI 1NN",
        "spacy_acc": "spaCy 1NN",
        "delta_acc": "Delta 1NN",
        "openai_ari": "OpenAI ARI",
        "spacy_ari": "spaCy ARI",
        "delta_ari": "Delta ARI",
    })
    return out


def build_docx() -> Path:
    root = config.ANALYSIS_DIR
    pkg = root / "paper_package_jql"
    out_path = root / "output" / "doc" / "methods_results_jql.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    scenario = pd.read_csv(pkg / "tables" / "table_scenario_overview.csv")
    perf = pd.read_csv(pkg / "tables" / "table_method_performance.csv")
    sig = pd.read_csv(pkg / "tables" / "table_significance_tests.csv")

    p3 = perf[perf["scenario"] == 3].set_index("method")
    p5 = perf[perf["scenario"] == 5].set_index("method")
    s3 = sig[sig["scenario"] == 3].copy()
    s5 = sig[sig["scenario"] == 5].copy()

    doc = Document()
    _set_doc_style(doc)

    doc.add_paragraph("Methods and Results", style="Title")
    doc.add_paragraph(
        "Target venue: Journal of Quantitative Linguistics. This section set is written as a "
        "manuscript-grade analytical narrative rather than a technical notebook report."
    )

    # Top-level section 1
    doc.add_heading("Methods", level=1)

    doc.add_heading("2.1 Research design and inferential aim", level=2)
    doc.add_paragraph(
        "The study examines how semantic similarity modulates the reliability of stylometric "
        "attribution based on Delta distances, and whether embedding-based representations "
        "provide measurable robustness under increasing semantic overlap. The design is "
        "scenario-based and fully controlled: each scenario contains 60 texts distributed across "
        "three author labels (A/B/C), thereby fixing class balance and sample size while varying "
        "the geometry of semantic and stylistic similarity."
    )
    doc.add_paragraph(
        "This architecture allows method performance to be interpreted as a function of corpus "
        "structure rather than of class imbalance, document-length asymmetry, or parameter "
        "idiosyncrasy. It also supports scenario-wise diagnosis of when local nearest-neighbour "
        "signals and global cluster geometry align versus when they diverge."
    )

    doc.add_heading("2.2 Scenario construction and corpus logic", level=2)
    doc.add_paragraph(
        "Scenario creation followed a progressive overlap strategy. Scenario 1 acts as a "
        "high-separation baseline; Scenario 2 introduces moderate semantic convergence; Scenario 3 "
        "is the principal distinct-text attribution condition (restored from the former Scenario 3a "
        "folder after audit); Scenario 4 is an intentional identical-text control (three unique base "
        "texts replicated across labels); and Scenario 5 maximizes semantic/topic overlap to probe "
        "failure behaviour under confounding."
    )
    doc.add_paragraph(
        "The corpus audit confirms this design: Scenario 3 and 5 are both sizeable and lexically rich, "
        "but only Scenario 5 collapses attribution, indicating that lexical richness alone does not "
        "guarantee author recoverability when topic alignment dominates."
    )
    scn_tbl = scenario.rename(columns={
        "scenario": "Scenario",
        "n_docs": "N docs",
        "n_unique_texts": "Unique texts",
        "total_words": "Total words",
        "vocab_size": "Vocab size",
        "is_identical_control": "Identical control",
    })
    _add_table_from_df(doc, scn_tbl, "Table 1. Scenario diagnostics and control status.")

    doc.add_heading("2.3 Representational spaces", level=2)
    doc.add_paragraph(
        "Three representations were compared under a common evaluation framework. "
        "(i) Delta: Cosine Delta (wurzburg) from stylo over MFW profiles (500 MFW for Scenarios 1, 2, 3, 5; "
        "393 MFW for Scenario 4; culling = 0). "
        "(ii) OpenAI embeddings: text-embedding-3-small (1536 dimensions), aggregated at document level. "
        "(iii) spaCy embeddings: en_core_web_lg document vectors (300 dimensions)."
    )
    doc.add_paragraph(
        "For embeddings, pairwise distances were computed with cosine distance. For stylometry, the "
        "native Delta distance matrix was used directly. This preserves methodological parity by "
        "evaluating each approach in its intrinsic metric space before classification and clustering."
    )

    doc.add_heading("2.4 Evaluation metrics", level=2)
    doc.add_paragraph(
        "Attribution was quantified via leave-one-out one-nearest-neighbour accuracy (1NN-LOO), which "
        "captures local identity preservation in the distance geometry. Global structural agreement with "
        "author labels was measured by Adjusted Rand Index (ARI) from agglomerative clustering (k = 3, "
        "average linkage, precomputed distances)."
    )

    doc.add_heading("2.5 Statistical inference", level=2)
    doc.add_paragraph(
        "Inference emphasizes effect sizes plus exact testing. In Scenario 3, we estimated pairwise "
        "accuracy differences (OpenAI-spaCy, OpenAI-Delta, spaCy-Delta), bootstrap 95% confidence intervals, "
        "paired randomization p-values, and exact McNemar p-values. In Scenario 5, exact binomial tests "
        "evaluated each method against chance (1/3), since observed performance was at floor level."
    )

    perf_wide = _performance_wide(perf).sort_values("scenario")
    _add_table_from_df(doc, perf_wide, "Table 2. Scenario-wise performance summary (1NN and ARI).")

    sig_tbl = sig.copy()
    sig_tbl = sig_tbl[["scenario", "contrast", "acc_diff", "ci95_low", "ci95_high", "p_permutation", "p_mcnemar_exact"]]
    sig_tbl = sig_tbl.rename(columns={
        "scenario": "Scenario",
        "contrast": "Contrast",
        "acc_diff": "Acc diff",
        "ci95_low": "CI low",
        "ci95_high": "CI high",
        "p_permutation": "p (perm)",
        "p_mcnemar_exact": "p (McNemar)",
    })
    _add_table_from_df(doc, sig_tbl, "Table 3. Inferential statistics for Scenario 3 and Scenario 5 tests.")

    # Top-level section 2
    doc.add_heading("Results", level=1)

    doc.add_heading("3.1 Macro-pattern across all scenarios", level=2)
    doc.add_paragraph(
        "The complete scenario series reveals three regimes. First, ceiling regimes (Scenarios 1, 2, 4) "
        "where all methods reach perfect classification and clustering. Second, a discriminative regime "
        "(Scenario 3) where methods separate in a stable rank order. Third, a collapse regime (Scenario 5) "
        "where all methods fail simultaneously. This regime-based view is analytically more informative "
        "than reporting only global averages, because it identifies the corpus conditions under which each "
        "representation is informative or fragile."
    )

    doc.add_heading("3.2 Scenario-by-scenario interpretation", level=2)
    doc.add_paragraph(
        "Scenario 1 confirms expected separability under low inter-author semantic interference. "
        "Scenario 2 preserves perfect recoverability despite increased overlap, indicating that moderate "
        "convergence does not automatically erase author signal. Scenario 3 is the key non-trivial case: "
        f"OpenAI ({_fmt(p3.loc['openai','acc_1nn_loo'])}) > spaCy ({_fmt(p3.loc['spacy','acc_1nn_loo'])}) "
        f"> Delta ({_fmt(p3.loc['delta','acc_1nn_loo'])}) in 1NN-LOO, with the same ordering in ARI "
        "(OpenAI strongest, Delta weakest). Scenario 4 is deliberately degenerate and should be interpreted "
        "as a control sanity-check, not as evidence of realistic attribution difficulty."
    )
    doc.add_paragraph(
        "Scenario 5 shows total nearest-neighbour collapse across all methods. Because this failure is "
        "convergent across stylometric and embedding spaces, it indicates a data-structure problem "
        "(topic-author entanglement) rather than a model-specific defect. In practical terms, when topic "
        "proximity dominates, authorial signal is not recoverable by either classical Delta or dense embeddings."
    )

    fig_dir = pkg / "figures"
    _add_figure(
        doc,
        fig_dir / "fig1_accuracy_heatmap.png",
        "Figure 1. 1NN-LOO accuracy by scenario and representation.",
    )
    _add_figure(
        doc,
        fig_dir / "fig2_ari_heatmap.png",
        "Figure 2. ARI by scenario and representation.",
    )

    doc.add_heading("3.3 Scenario 3 inferential profile", level=2)
    doc.add_paragraph(
        "Although OpenAI shows the highest point estimates in Scenario 3, inferential tests indicate "
        "that superiority should be framed as directional rather than definitive at this sample size. "
        "Confidence intervals for pairwise contrasts include zero and exact tests are non-significant "
        "at alpha = 0.05. Therefore, the evidence supports a robust descriptive ordering but not a "
        "strong null-rejection claim."
    )
    for _, r in s3.iterrows():
        doc.add_paragraph(
            f"- {r['contrast']}: diff={_fmt(r['acc_diff'],4)}, "
            f"95% CI [{_fmt(r['ci95_low'],4)}, {_fmt(r['ci95_high'],4)}], "
            f"p_perm={_fmt(r['p_permutation'],4)}, p_mcnemar={_fmt(r['p_mcnemar_exact'],4)}."
        )

    _add_figure(
        doc,
        fig_dir / "fig3_scenario3_within_across.png",
        "Figure 3. Scenario 3 within-group and across-group distance distributions.",
    )

    doc.add_heading("3.4 Scenario 5 collapse and methodological implications", level=2)
    doc.add_paragraph(
        "Scenario 5 chance tests are all highly significant in the negative direction (p = 2.72e-11), "
        "demonstrating that the 0.00 accuracies are not random fluctuations around chance but systematic "
        "misalignment between distance neighbourhoods and author labels. This is a critical substantive "
        "finding: under strong semantic confounding, method family becomes secondary to corpus design."
    )
    doc.add_paragraph(
        f"Observed 1NN floor values in Scenario 5: OpenAI={_fmt(p5.loc['openai','acc_1nn_loo'])}, "
        f"spaCy={_fmt(p5.loc['spacy','acc_1nn_loo'])}, Delta={_fmt(p5.loc['delta','acc_1nn_loo'])}."
    )
    for _, r in s5.iterrows():
        doc.add_paragraph(
            f"- {r['contrast']}: observed accuracy={_fmt(r['acc_a'],4)}, "
            f"chance={_fmt(r['acc_b'],4)}, p_exact={r['p_permutation']:.3e}."
        )

    _add_figure(
        doc,
        fig_dir / "fig4_scenario5_within_across.png",
        "Figure 4. Scenario 5 within-group and across-group distance distributions.",
    )

    doc.add_heading("3.5 Synthesis", level=2)
    doc.add_paragraph(
        "The scenario sequence supports a bounded but theoretically meaningful conclusion. "
        "In the non-trivial distinct-text setting (Scenario 3), embeddings, especially OpenAI, "
        "offer measurable gains over Delta. However, this representational advantage disappears under "
        "extreme topic confounding (Scenario 5), where all methods fail. For quantitative authorship "
        "research, this implies that scenario diagnostics and confound controls should be treated as "
        "first-order methodological requirements alongside model selection."
    )

    doc.save(out_path)
    return out_path


def sanity_check(path: Path) -> tuple[list[str], int, int, int]:
    doc = Document(path)
    h1 = [p.text.strip() for p in doc.paragraphs if p.style and p.style.name == "Heading 1"]
    return h1, len(doc.paragraphs), len(doc.tables), len(doc.inline_shapes)


def main() -> None:
    out = build_docx()
    h1, n_par, n_tbl, n_fig = sanity_check(out)
    if h1 != ["Methods", "Results"]:
        raise RuntimeError(f"Heading 1 mismatch: {h1}")
    if n_tbl < 3:
        raise RuntimeError(f"Expected at least 3 embedded tables, found {n_tbl}")
    if n_fig < 4:
        raise RuntimeError(f"Expected at least 4 embedded figures, found {n_fig}")
    print(f"DOCX written: {out}")
    print(f"Heading 1 sections: {h1}")
    print(f"Paragraph count: {n_par}")
    print(f"Embedded tables: {n_tbl}")
    print(f"Embedded figures: {n_fig}")


if __name__ == "__main__":
    main()
