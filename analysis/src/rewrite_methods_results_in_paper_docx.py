from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document


ROOT = Path(r"C:\Users\USER01\Dropbox\Workplace\D\George\PAPERS\Qualico 2025\Mikros, Cech")
PAPER_DOCX = ROOT / "Paper" / "260527_mikros_etal_semantic_smililarity.docx"
OUT_DOCX = ROOT / "Paper" / "260527_mikros_etal_semantic_smililarity_methods_results_rewritten.docx"
PKG = ROOT / "analysis" / "paper_package_jql" / "tables"


def _remove_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def _clean_tail_from_methods(doc: Document) -> None:
    """If Methods already exists, remove everything from it onward."""
    start_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip().lower()
        if t.startswith("2. methods") or t == "methods":
            start_idx = i
            break
    if start_idx is None:
        return
    for p in list(doc.paragraphs[start_idx:]):
        _remove_paragraph(p)


def _add_table(doc: Document, title: str, df: pd.DataFrame) -> None:
    doc.add_paragraph(title)
    t = doc.add_table(rows=1, cols=len(df.columns))
    try:
        t.style = "Table Grid"
    except KeyError:
        # Some templates rename/remove this style.
        pass
    for i, c in enumerate(df.columns):
        t.rows[0].cells[i].text = str(c)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, v in enumerate(row.tolist()):
            if isinstance(v, float):
                cells[i].text = f"{v:.4f}"
            else:
                cells[i].text = str(v)
    doc.add_paragraph("")


def _add_bullet(doc: Document, text: str) -> None:
    try:
        doc.add_paragraph(text, style="List Bullet")
    except KeyError:
        doc.add_paragraph(f"- {text}")


def main() -> None:
    scenario = pd.read_csv(PKG / "table_scenario_overview.csv")
    perf = pd.read_csv(PKG / "table_method_performance.csv")
    sig = pd.read_csv(PKG / "table_significance_tests.csv")

    # Compact tables for in-text insertion
    scenario_t = scenario.copy()
    perf_t = perf.copy()
    perf_t["acc_1nn_loo"] = perf_t["acc_1nn_loo"].round(4)
    perf_t["ari_3cluster"] = perf_t["ari_3cluster"].round(4)
    sig_t = sig.copy()
    for c in ("acc_a", "acc_b", "acc_diff", "ci95_low", "ci95_high", "p_permutation", "p_mcnemar_exact"):
        if c in sig_t.columns:
            sig_t[c] = sig_t[c].round(6)

    p3 = perf[perf["scenario"] == 3].set_index("method")
    p5 = perf[perf["scenario"] == 5].set_index("method")

    doc = Document(PAPER_DOCX)
    _clean_tail_from_methods(doc)

    # ------------------------------------------------------------------
    # 2. METHODS
    # ------------------------------------------------------------------
    doc.add_paragraph("2. Methods")

    doc.add_paragraph("2.1 Research design and rationale")
    doc.add_paragraph(
        "To isolate the effect of semantic similarity on stylometric attribution, we "
        "used a controlled scenario-based design. Rather than evaluating one corpus "
        "in isolation, we constructed a sequence of experimental conditions that vary "
        "in lexical-semantic overlap while keeping class cardinality constant "
        "(three author labels: A, B, C; 60 texts per scenario). This design allows "
        "method behaviour to be interpreted as a function of corpus geometry, not as "
        "an artifact of unequal sample sizes."
    )
    doc.add_paragraph(
        "The central methodological decision was to compare a classical distance-based "
        "stylometric representation (Cosine Delta) against two embedding spaces "
        "(OpenAI text-embedding-3-small and spaCy en_core_web_lg) under identical "
        "classification and clustering protocols. The comparison therefore targets "
        "the representational layer, not classifier engineering."
    )

    doc.add_paragraph("2.2 Scenario creation workflow")
    doc.add_paragraph(
        "The scenario series was generated as controlled artificial corpora, with "
        "each scenario encoding a different balance between within-author coherence "
        "and between-author similarity. Empirically, scenario construction can be "
        "recovered from corpus diagnostics (number of unique texts, vocabulary size) "
        "and from distance-separation behaviour observed in pilot analyses."
    )
    doc.add_paragraph(
        "Scenario 1 serves as a high-separation baseline (60/60 unique texts; "
        "strong between-author gap). Scenario 2 introduces moderate semantic "
        "convergence while preserving label recoverability. Scenario 3 (restored from "
        "the folder previously named Scenario 3a) is the principal distinct-text test "
        "condition and was the source of the original saved Scenario 3 distance table. "
        "Scenario 4 is an intentional identical-text control (three unique base texts "
        "replicated across labels), designed to verify pipeline behaviour in a "
        "degenerate but diagnostically informative setting. Scenario 5 introduces the "
        "highest level of inter-author semantic overlap, producing topic-author "
        "entanglement and strong attribution difficulty."
    )
    _add_table(doc, "Table 1. Scenario diagnostics used to interpret scenario construction.", scenario_t)

    doc.add_paragraph("2.3 Text representations")
    doc.add_paragraph(
        "We evaluated three representations. (i) Cosine Delta distances were computed "
        "with stylo on Most Frequent Word profiles (500 MFW in Scenarios 1, 2, 3, and 5; "
        "393 MFW in Scenario 4 due to available feature constraints; culling = 0). "
        "(ii) OpenAI embeddings used text-embedding-3-small (1536 dimensions) with "
        "document-level averaging. (iii) spaCy embeddings used en_core_web_lg "
        "(300 dimensions) and document vectors derived from the complete token stream."
    )
    doc.add_paragraph(
        "Embedding distances were cosine-based. For stylometry, the distance matrix "
        "was taken directly from Cosine Delta output. This ensures that each method "
        "is evaluated in its native representational space before downstream inference."
    )

    doc.add_paragraph("2.4 Evaluation protocol")
    doc.add_paragraph(
        "Primary attribution performance was measured with leave-one-out 1-nearest-"
        "neighbour accuracy (1NN-LOO). To assess whether global partition structure "
        "matches author labels, we additionally computed Adjusted Rand Index (ARI) "
        "from agglomerative clustering (k = 3, average linkage, precomputed distances). "
        "Using both local (1NN) and global (ARI) criteria helps distinguish cases where "
        "nearest-neighbour behaviour and cluster geometry diverge."
    )

    doc.add_paragraph("2.5 Inference and uncertainty")
    doc.add_paragraph(
        "For Scenario 3, we tested pairwise method contrasts using bootstrap confidence "
        "intervals of accuracy differences, paired randomization tests, and exact "
        "McNemar tests. For Scenario 5, exact binomial tests compared observed "
        "accuracy against a chance baseline (1/3). This multi-layer inferential "
        "strategy supports effect-size interpretation alongside exact hypothesis testing."
    )
    _add_table(doc, "Table 2. Method-level performance by scenario.", perf_t)
    _add_table(doc, "Table 3. Scenario 3 contrasts and Scenario 5 chance tests.", sig_t)

    # ------------------------------------------------------------------
    # 3. RESULTS
    # ------------------------------------------------------------------
    doc.add_paragraph("3. Results")

    doc.add_paragraph("3.1 Cross-scenario performance profile")
    doc.add_paragraph(
        "The five-scenario series reveals a non-linear relation between semantic "
        "similarity and attribution reliability. In Scenarios 1, 2, and 4, all methods "
        "reach 1.00 1NN-LOO and 1.00 ARI, indicating either strongly separable or "
        "control-like geometry. In Scenario 3, methods diverge in expected order: "
        f"OpenAI ({p3.loc['openai','acc_1nn_loo']:.4f}) > "
        f"spaCy ({p3.loc['spacy','acc_1nn_loo']:.4f}) > "
        f"Delta ({p3.loc['delta','acc_1nn_loo']:.4f}). "
        "In Scenario 5, all methods collapse to zero accuracy, indicating a "
        "structural failure mode rather than random fluctuation."
    )

    doc.add_paragraph("3.2 Scenario-specific interpretation")
    doc.add_paragraph(
        "Scenario 1 confirms that all representations can recover authorship when "
        "between-author signal is high relative to within-author variability. Scenario 2 "
        "shows that moderate semantic convergence is still tractable for all methods, "
        "suggesting that stylometric and embedding spaces remain aligned under partial overlap."
    )
    doc.add_paragraph(
        "Scenario 3 is the key inferential condition because it is neither trivial "
        "nor degenerate. Here, OpenAI embeddings provide the highest local attribution "
        "accuracy and best global partition recovery, while spaCy remains competitive "
        "and Delta trails both embedding methods. This pattern supports the claim that "
        "semantic representation contributes additional discriminative information when "
        "topic overlap is present but not dominant."
    )
    doc.add_paragraph(
        "Scenario 4 should not be interpreted as ordinary success despite perfect scores: "
        "it is an identical-text control with only three unique base texts. Its function "
        "is diagnostic (pipeline sanity and metric behaviour), not evidential about real-world "
        "authorship complexity."
    )
    doc.add_paragraph(
        "Scenario 5 exhibits complete nearest-neighbour failure for every method "
        f"(OpenAI {p5.loc['openai','acc_1nn_loo']:.4f}; "
        f"spaCy {p5.loc['spacy','acc_1nn_loo']:.4f}; "
        f"Delta {p5.loc['delta','acc_1nn_loo']:.4f}). "
        "Because all three representations fail simultaneously, the most plausible "
        "interpretation is topic confounding: pairwise proximity primarily reflects "
        "shared semantic content across author labels, overwhelming author-specific style."
    )

    doc.add_paragraph("3.3 Scenario 3 inferential contrasts")
    doc.add_paragraph(
        "Pairwise contrasts in Scenario 3 show positive effect directions for OpenAI "
        "against both alternatives, but uncertainty intervals overlap zero and exact "
        "tests are not significant at alpha = 0.05. The empirical ranking "
        "OpenAI > spaCy > Delta is therefore robust as a descriptive ordering, while "
        "inferentially it should be treated as directional evidence that invites "
        "confirmation on larger or more heterogeneous corpora."
    )
    for _, row in sig[sig["scenario"] == 3].iterrows():
        _add_bullet(
            doc,
            f"{row['contrast']}: acc diff={row['acc_diff']:.4f}, "
            f"95% CI [{row['ci95_low']:.4f}, {row['ci95_high']:.4f}], "
            f"p_perm={row['p_permutation']:.4f}, "
            f"p_mcnemar={row['p_mcnemar_exact']:.4f}.",
        )

    doc.add_paragraph("3.4 Implications for stylometric methodology")
    doc.add_paragraph(
        "Taken together, the scenario series supports a conditional conclusion: "
        "embedding-enhanced representations can outperform classical Delta in non-trivial "
        "semantic-overlap settings (Scenario 3), but no representation is immune to "
        "design-level confounds (Scenario 5). For quantitative stylometry, this implies "
        "that evaluation should explicitly model scenario geometry and include control "
        "conditions, not only aggregate accuracy scores."
    )

    doc.save(OUT_DOCX)
    print(str(OUT_DOCX))


if __name__ == "__main__":
    main()
